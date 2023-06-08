from docx import Document
from docx.shared import Inches


class WordWriter:
    def __init__(self, template, result_path):
        self.template_doc = Document(template)
        self.result_path = result_path

    def _get_cell(self, mark):
        for table in self.template_doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip() == mark:
                        return cell

    def set_photo(self, img_path):
        cell = self._get_cell("photo")
        if cell is not None:
            self._add_cell_image(img_path, cell)


    def _add_cell_image__(self, img_path, cell):
        cell.text = ''
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run()
        run.add_picture(img_path, width=Inches(1.8), height=Inches(2.2))

    def set_name(self, name):
        cell = self._get_cell("name")
        if cell is not None:
            cell.text = name

    def set_birth(self, birth):
        cell = self._get_cell("birth")
        if cell is not None:
            cell.text = str(birth)

    def set_company(self, company):
        cell = self._get_cell("company")
        if cell is not None:
            cell.text = company

    def set_work_code(self, work_code):
        cell = self._get_cell("work_code")
        if cell is not None:
            cell.text = work_code

    def set_address(self, address):
        cell = self._get_cell("address")
        if cell is not None:
            cell.text = address

    def set_phone(self, phone):
        cell = self._get_cell("phone")
        if cell is not None:
            cell.text = str(phone)

    def set_id(self, id):
        cell = self._get_cell("id")
        if cell is not None:
            cell.text = str(id)

    def set_sex(self, sex):
        cell = self._get_cell("sex")
        if cell is not None:
            cell.text = sex

    def save(self):
        self.template_doc.save(self.result_path)
